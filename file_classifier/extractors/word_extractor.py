import io
import base64
import zipfile
import structlog
from docx import Document
from PIL import Image
from file_classifier.extractors.base import BaseExtractor, ExtractionResult

logger = structlog.get_logger()

MAX_PARAGRAPHS = 200
MIN_TEXT_LENGTH = 50
MAX_IMAGE_DIMENSION = 2048


class WordExtractor(BaseExtractor):
    def extract(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        logger.info("Extracting Word content", filename=filename)
        buf = io.BytesIO(file_bytes)
        doc = Document(buf)

        parts = []

        para_count = 0
        for para in doc.paragraphs:
            if para_count >= MAX_PARAGRAPHS:
                break
            text = para.text.strip()
            if not text:
                continue

            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)
            para_count += 1

        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            if rows:
                header = rows[0]
                table_str = " | ".join(header) + "\n"
                table_str += " | ".join("---" for _ in header) + "\n"
                for row in rows[1:20]:
                    table_str += " | ".join(row) + "\n"
                parts.append(f"\n### Table {t_idx + 1}\n{table_str}")

        text = "\n\n".join(parts)

        if len(text.strip()) < MIN_TEXT_LENGTH:
            logger.info("Word doc has minimal text, attempting image extraction", filename=filename)
            image_b64 = self._extract_first_embedded_image(file_bytes)
            if image_b64:
                return ExtractionResult(
                    text_content="[Word document is image-based - content sent as image for visual analysis]",
                    metadata={
                        "paragraphs_extracted": para_count,
                        "tables_found": len(doc.tables),
                        "extraction_method": "image",
                    },
                    is_image_based=True,
                    image_base64=image_b64,
                )

        return ExtractionResult(
            text_content=text,
            metadata={"paragraphs_extracted": para_count, "tables_found": len(doc.tables)},
        )

    def _extract_first_embedded_image(self, file_bytes: bytes) -> str | None:
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                image_names = sorted(
                    [n for n in z.namelist() if n.startswith("word/media/")]
                )
                if not image_names:
                    return None

                largest = max(image_names, key=lambda n: z.getinfo(n).file_size)
                with z.open(largest) as f:
                    img_bytes = f.read()

            img = Image.open(io.BytesIO(img_bytes))
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")
            if max(img.size) > MAX_IMAGE_DIMENSION:
                img.thumbnail((MAX_IMAGE_DIMENSION, MAX_IMAGE_DIMENSION), Image.LANCZOS)

            out = io.BytesIO()
            img.save(out, format="PNG")
            return base64.b64encode(out.getvalue()).decode("utf-8")
        except Exception as e:
            logger.warning("Failed to extract embedded image from docx", error=str(e))
            return None
